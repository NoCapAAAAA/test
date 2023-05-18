from core import models as m
from django.contrib import messages
from service.filters import UsersFilterDirector
from django.contrib.auth.models import Group
from organization import forms as f
from django.utils.decorators import method_decorator
from django.views.generic import UpdateView,ListView
from django.urls import reverse_lazy
from service.decorators import group_required
from django.contrib.auth import get_user_model
from django.db.models import Sum
from django.views.generic import TemplateView
from django.http import HttpResponse
from docx import *
from django.utils import timezone
from io import BytesIO
from datetime import datetime, date
import locale
from django.db.models import Count
from django.db.models.functions import ExtractYear, ExtractMonth
from django.http import JsonResponse
from core.models import OrderStorage
from service.charts import months, colorPrimary, colorSuccess, colorDanger, generate_color_palette, get_year_dict
from authentication.models import User as Users
User = get_user_model()


class DirectorHomeView(TemplateView):
    @method_decorator(group_required('Директор'))
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    template_name = 'director/home_director.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        last_week = timezone.now() - timezone.timedelta(days=7)
        context['orders'] = OrderStorage.objects.count()
        context['lastweakusers'] = User.objects.filter(date_joined__gte=last_week).count()
        context['lastweakorders'] = OrderStorage.objects.filter(created_at__gte=last_week).count()
        context['instorageorders'] = OrderStorage.objects.filter(status=m.OrderStatus.STORAGE).count()
        context['finishorders'] = OrderStorage.objects.filter(status=m.OrderStatus.FINISH).count()
        last_month = timezone.now() - timezone.timedelta(days=30)
        context['totalpriceorderlastmonth'] = OrderStorage.objects.filter(
            is_payed=True,
            status=m.OrderStatus.FINISH,
            created_at__gte=last_month
        ).aggregate(Sum('price'))['price__sum']
        context['totalpriceorderyear'] = OrderStorage.objects.filter(is_payed=True, status=m.OrderStatus.FINISH).aggregate(Sum('price'))['price__sum']
        last_year = timezone.now() - timezone.timedelta(days=365)
        context['totalpriceorderlastyear'] = OrderStorage.objects.filter(
            is_payed=True,
            status=m.OrderStatus.FINISH,
            created_at__year=last_year.year - 1
        ).aggregate(Sum('price'))['price__sum']
        context['callaplication'] = m.CallApplication.objects.order_by('-pk')[:5]
        return context




class DirectorUsersListView(ListView):
    @method_decorator(group_required('Директор'))
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    template_name = 'director/users_list_director.html'
    model = User
    paginate_by = 7

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['users'] = User.objects.all()
        context['filter'] = UsersFilterDirector(self.request.GET, queryset=self.get_queryset())
        return context

    def get_queryset(self, **kwargs):
        search_results = UsersFilterDirector(self.request.GET, self.queryset)
        self.no_search_result = True if not search_results.qs else False
        return search_results.qs.distinct()


class DirectorUpdateUserView(UpdateView):
    @method_decorator(group_required('Директор'))
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    template_name = 'director/user_get_permissions.html'
    model = User
    form_class = f.SettingsProfile

    def get_queryset(self):
        return User.objects.filter(pk=self.kwargs['pk'])

    def form_valid(self, form):
        messages.success(self.request, "The user was updated successfully.")
        return super(DirectorUpdateUserView, self).form_valid(form)

    def get_success_url(self):
        return reverse_lazy('director_get_user_permissions', kwargs={'pk': self.kwargs['pk']})


class DirectorEmployeeDetailView(UpdateView):
    @method_decorator(group_required('Директор'))
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    template_name = 'director/detail_employee_director.html'
    model = User
    form_class = f.SettingsProfile

    def form_valid(self, form):
        user = form.save(commit=False)
        user.groups.clear()  # очищаем группы пользователя
        groups = self.request.POST.getlist('groups')  # получаем список выбранных групп
        for group_id in groups:
            group = Group.objects.get(id=group_id)
            user.groups.add(group)  # добавляем пользователя в группы
        return super().form_valid(form)

    def post(self, request, *args, **kwargs):
        self.object = self.get_object()
        group_id = request.POST.get('group_id')
        if group_id:
            group = Group.objects.get(id=group_id)
            self.object.groups.remove(group)
        return super().post(request, *args, **kwargs)

    def get_queryset(self):
        return User.objects.filter(pk=self.kwargs['pk'])

    def get_success_url(self):
        return reverse_lazy('director_detail_employee_view', kwargs={'pk': self.kwargs['pk']})


class DirectorListEmployeeView(ListView):
    @method_decorator(group_required('Директор'))
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    template_name = 'director/list_employee_director.html'
    model = User
    paginate_by = 7
    context_object_name = 'employees'
    queryset = User.objects.filter(groups__name='Менеджер')


class DirectorUsersReportView(TemplateView):
    @method_decorator(group_required('Директор'))
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    template_name = 'director/users_report_director.html'


def create_report(request):
    locale.setlocale(
        category=locale.LC_ALL,
        locale="Russian"  # Note: do not use "de_DE" as it doesn't work
    )
    current_datetime = datetime.now()
    str_current_datetime = str(current_datetime)
    document = Document()
    docx_title = "report" + str_current_datetime + ".docx"
    # ---- Cover Letter ----
    document.add_paragraph()
    document.add_paragraph("%s" % date.today().strftime('%B %d, %Y'))

    document.add_paragraph('Отчёт о заказах')
    orders = m.OrderStorage.objects.all().order_by('pk')
    orders_count = orders.count()
    table = document.add_table(rows=orders_count + 1, cols=8)
    table.style = 'Table Grid'
    table.cell(0, 0).text = 'Номер заказа'
    table.cell(0, 1).text = 'Клиент'
    table.cell(0, 2).text = 'Размер шины'
    table.cell(0, 3).text = 'Период хранения'
    table.cell(0, 4).text = 'Адрес сервиса'
    table.cell(0, 5).text = 'Статус заказа'
    table.cell(0, 6).text = 'Стоимость заказа'
    table.cell(0, 7).text = 'Оплачен'

    # Adding data to the table
    for i, order in enumerate(orders, start=1):
        table.cell(i, 0).text = str(order.pk)
        table.cell(i, 1).text = str(order.user)
        table.cell(i, 2).text = str(order.size)
        table.cell(i, 3).text = str(f"{str(order.period)} мес.")
        table.cell(i, 4).text = str(order.adress) if order.adress else "---"
        table.cell(i, 5).text = str(order.get_status_display())
        table.cell(i, 6).text = str(f"{str(order.price)} ₽")
        table.cell(i, 7).text = "Да" if order.is_payed else "Нет"

    document.add_page_break()

    # Prepare document for download
    # -----------------------------
    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length
    return response


def create_report_users(request):
    locale.setlocale(
        category=locale.LC_ALL,
        locale="Russian"  # Note: do not use "de_DE" as it doesn't work
    )
    current_datetime = datetime.now()
    str_current_datetime = str(current_datetime)
    document = Document()
    docx_title = "report" + str_current_datetime + ".docx"
    # ---- Cover Letter ----
    document.add_paragraph()
    document.add_paragraph("%s" % date.today().strftime('%B %d, %Y'))

    document.add_paragraph('Отчёт о пользователях')

    users_count = Users.objects.all().order_by('pk').count()

    table = document.add_table(rows=users_count + 1, cols=8)
    table.style = 'Table Grid'
    table.cell(0, 0).text = 'ID'
    table.cell(0, 1).text = 'Имя'
    table.cell(0, 2).text = 'Фамилия'
    table.cell(0, 3).text = 'Отчество'
    table.cell(0, 4).text = 'Телефон'
    table.cell(0, 5).text = 'Email'
    table.cell(0, 6).text = 'Пол'

    # Adding data to the table
    for i, user in enumerate(Users.objects.all(), start=1):
        table.cell(i, 0).text = str(user.pk)
        table.cell(i, 1).text = str(user.first_name)
        table.cell(i, 2).text = str(user.last_name)
        table.cell(i, 3).text = str(user.middle_name)
        table.cell(i, 4).text = str(user.phone_number)
        table.cell(i, 5).text = str(user.email)
        table.cell(i, 6).text = str(user.gender)

    document.add_page_break()

    # Prepare document for download
    # -----------------------------
    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length
    return response

def get_filter_options(request):
    grouped_purchases = OrderStorage.objects.annotate(year=ExtractYear("created_at")).values("year").order_by("-year").distinct()
    options = [purchase["year"] for purchase in grouped_purchases]

    return JsonResponse({
        "options": options,
    })


def get_sales_chart(request, year):
    purchases = OrderStorage.objects.filter(created_at__year=year)
    grouped_purchases = purchases.annotate(month=ExtractMonth("created_at"))\
        .values("month").annotate(count=Count("id")).values("month", "count").order_by("month")

    sales_dict = get_year_dict()

    for group in grouped_purchases:
        sales_dict[months[group["month"]-1]] = group["count"]

    return JsonResponse({
        "title": f"Sales in {year}",
        "data": {
            "labels": list(sales_dict.keys()),
            "datasets": [{
                "label": "Количество заказов",
                "backgroundColor": colorPrimary,
                "borderColor": colorPrimary,
                "data": list(sales_dict.values()),
            }]
        },
    })
